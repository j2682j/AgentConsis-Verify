from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from tools.attachment_reader import (
    AttachmentEvidenceBuilder,
    AttachmentPayloadBuilder,
    AttachmentProfileBuilder,
)
from tools.attachment_strategy import AttachmentStrategy, AttachmentStrategyExecutor
from tools.deterministic_handlers import DeterministicHandlerRouter


def _router_metadata(read_result: dict) -> dict:
    return {
        "attachment_profile": read_result["profile"],
        "parsed_payload": read_result["parsed_payload"],
        "require_attachment_provenance": True,
    }


class _CapabilityRecordingPlanner:
    def __init__(self, strategy: AttachmentStrategy) -> None:
        self.strategy = strategy
        self.allowed_handlers: list[dict] = []

    def plan(self, *, question, attachment_profile, allowed_handlers):
        del question, attachment_profile
        self.allowed_handlers = list(allowed_handlers)
        return self.strategy, '{"planned":true}'


class AttachmentHandlerCapabilityTests(unittest.TestCase):
    def test_xlsx_exposes_only_ready_attachment_handlers_and_runs_typed_rows(self):
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "records.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.append(["name", "value"])
            sheet.append(["A", 10])
            sheet.append(["B", 20])
            sheet["A2"].fill = PatternFill(fill_type="solid", fgColor="FFFF00")
            workbook.save(path)
            read_result = AttachmentEvidenceBuilder().build(
                "How many records are in the table?",
                {"file_path": str(path), "extension": ".xlsx"},
            )

            router = DeterministicHandlerRouter()
            capabilities, _ = router.eligible_capabilities(
                question="How many records are in the table?",
                attachment={"file_path": str(path), "extension": ".xlsx"},
                attachment_result=read_result["context"],
                metadata=_router_metadata(read_result),
            )
            names = {item.handler_name for item in capabilities}
            result = router.run(
                question="How many records are in the table?",
                attachment={"file_path": str(path), "extension": ".xlsx"},
                attachment_result=read_result["context"],
                handler_name="table_exact_operations",
                metadata={
                    **_router_metadata(read_result),
                    "eligible_handler_names": sorted(names),
                },
            )

        self.assertEqual(names, {"table_exact_operations", "table_aggregation"})
        self.assertEqual(result.answer, "2")
        self.assertIn("cell_metadata", read_result["profile"]["available_inputs"])

    def test_strategy_planner_receives_capabilities_not_full_registry(self):
        from openpyxl import Workbook

        planner = _CapabilityRecordingPlanner(
            AttachmentStrategy(
                information_need="Count table records.",
                required_handler="table_exact_operations",
                required_inputs=["rows"],
                expected_answer="record count",
            )
        )
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "records.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.append(["name"])
            sheet.append(["A"])
            workbook.save(path)
            result = AttachmentStrategyExecutor(planner=planner).run(
                question="How many records are in the table?",
                attachment={"file_path": str(path), "extension": ".xlsx"},
            )

        names = {item["handler_name"] for item in planner.allowed_handlers}
        self.assertEqual(names, {"table_exact_operations", "table_aggregation"})
        self.assertEqual(result.handler_status, "success")

    def test_pdb_requires_two_explicit_targets_before_distance_handler_is_eligible(self):
        pdb_text = (
            "ATOM      1  N   ALA A   1       0.000   0.000   0.000  1.00 20.00           N\n"
            "ATOM      2  CA  ALA A   1       3.000   4.000   0.000  1.00 20.00           C\n"
        )
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.pdb"
            path.write_text(pdb_text, encoding="utf-8")
            builder = AttachmentEvidenceBuilder()
            ambiguous = builder.build(
                "Compute the distance between the selected atoms.",
                {"file_path": str(path), "extension": ".pdb"},
            )
            explicit = builder.build(
                "Compute the distance between atom 1 and atom 2.",
                {"file_path": str(path), "extension": ".pdb"},
            )
            router = DeterministicHandlerRouter()
            ambiguous_caps, _ = router.eligible_capabilities(
                question="Compute the distance between the selected atoms.",
                attachment={"file_path": str(path), "extension": ".pdb"},
                attachment_result=ambiguous["context"],
                metadata=_router_metadata(ambiguous),
            )
            explicit_caps, _ = router.eligible_capabilities(
                question="Compute the distance between atom 1 and atom 2.",
                attachment={"file_path": str(path), "extension": ".pdb"},
                attachment_result=explicit["context"],
                metadata=_router_metadata(explicit),
            )
            explicit_names = [item.handler_name for item in explicit_caps]
            result = router.run(
                question="Compute the distance between atom 1 and atom 2.",
                attachment={"file_path": str(path), "extension": ".pdb"},
                attachment_result=explicit["context"],
                handler_name="coordinate_distance",
                metadata={
                    **_router_metadata(explicit),
                    "eligible_handler_names": explicit_names,
                },
            )

        self.assertNotIn("coordinate_distance", {item.handler_name for item in ambiguous_caps})
        self.assertIn("coordinate_distance", explicit_names)
        self.assertEqual(result.answer, "5")

    def test_jsonld_relations_feed_graph_handler(self):
        data = (
            '{"@id":"A","next":{"@id":"B","next":{"@id":"C"}}}'
        )
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "graph.jsonld"
            path.write_text(data, encoding="utf-8")
            read_result = AttachmentEvidenceBuilder().build(
                'Find the shortest path from "A" to "C".',
                {"file_path": str(path), "extension": ".jsonld"},
            )
            router = DeterministicHandlerRouter()
            capabilities, _ = router.eligible_capabilities(
                question='Find the shortest path from "A" to "C".',
                attachment={"file_path": str(path), "extension": ".jsonld"},
                attachment_result=read_result["context"],
                metadata=_router_metadata(read_result),
            )
            names = [item.handler_name for item in capabilities]
            result = router.run(
                question='Find the shortest path from "A" to "C".',
                attachment={"file_path": str(path), "extension": ".jsonld"},
                attachment_result=read_result["context"],
                handler_name="graph_shortest_path",
                metadata={
                    **_router_metadata(read_result),
                    "eligible_handler_names": names,
                },
            )

        self.assertIn("graph_shortest_path", names)
        self.assertEqual(result.answer, "A -> B -> C")

    def test_visual_json_builds_typed_numbers_without_model_call(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "image.png"
            path.write_bytes(b"test")
            content = (
                '{"ocr_blocks":[{"text":"12 and 7","region":"center"}],'
                '"numbers":[12,7],"colors":["red"],"summary":"two values"}'
            )
            payload = AttachmentPayloadBuilder().build(
                file_path=path,
                extension=".png",
                content=content,
                reader="test_vision",
            )
            profile = AttachmentProfileBuilder().build(
                file_path=path,
                extension=".png",
                read_ok=True,
                reader="test_vision",
                content=content,
                parsed_payload=payload,
            )

        self.assertIn("numbers", profile.available_inputs)
        self.assertEqual(payload.visual_blocks[-1].attributes["numbers"], [12, 7])

    def test_docx_prose_does_not_expose_graph_edges(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "notes.docx"
            document = Document()
            document.add_heading("Notes", level=1)
            document.add_paragraph("Alpha, beta, and gamma are ordinary prose.")
            document.save(path)
            read_result = AttachmentEvidenceBuilder().build(
                "Summarize the document.",
                {"file_path": str(path), "extension": ".docx"},
            )

        self.assertNotIn("edges", read_result["profile"]["available_inputs"])
        self.assertEqual(read_result["parsed_payload"]["text_blocks"][0]["block_type"], "heading")


if __name__ == "__main__":
    unittest.main()
