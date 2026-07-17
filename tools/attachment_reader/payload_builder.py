from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Any

from .models import (
    CoordinateRecord,
    ListPayload,
    ParsedAttachmentPayload,
    RelationRecord,
    TablePayload,
    TextBlock,
    VisualBlock,
)


class AttachmentPayloadBuilder:
    """將不同附件格式轉成統一的 typed payload。"""

    def __init__(self, *, max_table_rows: int = 80) -> None:
        self.max_table_rows = max(1, int(max_table_rows))

    def build(
        self,
        *,
        file_path: Path,
        extension: str,
        content: str,
        reader: str,
        reader_metadata: dict[str, Any] | None = None,
    ) -> ParsedAttachmentPayload:
        payload = ParsedAttachmentPayload(
            native_metadata={
                **dict(reader_metadata or {}),
                "attachment_kind": self._attachment_kind(extension),
            },
            provenance={
                "source": "attachment_reader",
                "file_path": str(file_path),
                "file_type": extension,
                "reader": reader,
            },
        )
        if not file_path.exists() or not file_path.is_file():
            return payload

        try:
            if extension == ".csv":
                self._read_csv(file_path, payload)
            elif extension in {".xlsx", ".xls"}:
                self._read_workbook(file_path, extension, payload)
            elif extension == ".docx":
                self._read_docx(file_path, payload)
            elif extension == ".pdf":
                self._read_page_blocks(content, payload, marker="page")
            elif extension == ".pptx":
                self._read_page_blocks(content, payload, marker="slide")
            elif extension in {".json", ".jsonld"}:
                self._read_json(file_path, payload)
            elif extension == ".pdb":
                self._read_pdb(file_path, payload)
            elif extension in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}:
                self._read_visual_content(content, payload)
            elif content.strip():
                payload.text_blocks.append(TextBlock(text=content.strip()))
        except Exception as exc:
            payload.native_metadata["structured_parse_error"] = f"{type(exc).__name__}: {exc}"

        if not payload.has_content() and content.strip():
            payload.text_blocks.append(TextBlock(text=content.strip()))
        return payload

    @staticmethod
    def _attachment_kind(extension: str) -> str:
        if extension in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}:
            return "image"
        if extension in {".mp3", ".m4a", ".wav", ".flac", ".ogg"}:
            return "audio"
        if extension in {".mp4", ".mov", ".mkv", ".avi", ".webm"}:
            return "video"
        if extension == ".zip":
            return "archive"
        if extension in {".py", ".yaml", ".yml", ".xml", ".html"}:
            return "code"
        return "document"

    def from_text(
        self,
        *,
        content: str,
        provenance: dict[str, Any] | None = None,
    ) -> ParsedAttachmentPayload:
        payload = ParsedAttachmentPayload(provenance=dict(provenance or {}))
        if content.strip():
            payload.text_blocks.append(TextBlock(text=content.strip()))
        return payload

    def _read_csv(self, file_path: Path, payload: ParsedAttachmentPayload) -> None:
        with file_path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
            rows = [
                [str(cell) for cell in row]
                for _, row in zip(range(self.max_table_rows + 1), csv.reader(handle))
            ]
        if not rows:
            return
        columns = rows[0]
        payload.tables.append(
            TablePayload(
                name="CSV",
                columns=columns,
                rows=rows[1:],
                truncated=len(rows) >= self.max_table_rows + 1,
            )
        )

    def _read_workbook(
        self,
        file_path: Path,
        extension: str,
        payload: ParsedAttachmentPayload,
    ) -> None:
        if extension == ".xlsx":
            import openpyxl

            workbook = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            try:
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    rows: list[list[str]] = []
                    cell_metadata: dict[str, dict[str, Any]] = {}
                    for row_index, row in enumerate(sheet.iter_rows()):
                        if row_index >= self.max_table_rows + 1:
                            break
                        rows.append(["" if cell.value is None else str(cell.value) for cell in row])
                        for cell in row:
                            metadata = self._cell_metadata(cell)
                            if metadata:
                                cell_metadata[str(cell.coordinate)] = metadata
                    if rows:
                        payload.tables.append(
                            TablePayload(
                                name=sheet_name,
                                columns=rows[0],
                                rows=rows[1:],
                                cell_metadata=cell_metadata,
                                truncated=int(sheet.max_row or 0) > self.max_table_rows + 1,
                            )
                        )
            finally:
                workbook.close()
            return

        import pandas as pd

        sheets = pd.read_excel(file_path, sheet_name=None, dtype=object, engine="xlrd")
        for sheet_name, dataframe in sheets.items():
            preview = dataframe.head(self.max_table_rows)
            payload.tables.append(
                TablePayload(
                    name=str(sheet_name),
                    columns=[str(value) for value in preview.columns.tolist()],
                    rows=[
                        ["" if value is None else str(value) for value in row.tolist()]
                        for _, row in preview.iterrows()
                    ],
                    truncated=len(dataframe) > self.max_table_rows,
                )
            )

    def _read_docx(self, file_path: Path, payload: ParsedAttachmentPayload) -> None:
        from docx import Document

        document = Document(str(file_path))
        list_items: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if "list" in style_name.lower():
                list_items.append(text)
            else:
                block_type = "heading" if "heading" in style_name.lower() else "paragraph"
                payload.text_blocks.append(
                    TextBlock(text=text, section=style_name, block_type=block_type)
                )
        if list_items:
            payload.lists.append(ListPayload(title="document list", items=list_items))
        for index, table in enumerate(document.tables, start=1):
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows[: self.max_table_rows + 1]
            ]
            if rows:
                payload.tables.append(
                    TablePayload(
                        name=f"table {index}",
                        columns=rows[0],
                        rows=rows[1:],
                        truncated=len(table.rows) > self.max_table_rows + 1,
                    )
                )

    def _read_json(self, file_path: Path, payload: ParsedAttachmentPayload) -> None:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        data = json.loads(text)
        payload.text_blocks.append(
            TextBlock(text=json.dumps(data, ensure_ascii=False, indent=2))
        )
        self._json_relations(data, payload=payload, parent="root")

    def _json_relations(
        self,
        value: Any,
        *,
        payload: ParsedAttachmentPayload,
        parent: str,
    ) -> None:
        if len(payload.relations) >= 500:
            return
        if isinstance(value, dict):
            source = str(value.get("@id") or value.get("id") or value.get("name") or parent)
            for key, child in value.items():
                if key in {"@id", "id", "name"}:
                    continue
                if isinstance(child, (str, int, float, bool)):
                    payload.relations.append(
                        RelationRecord(source=source, relation=str(key), target=str(child))
                    )
                elif isinstance(child, dict):
                    target_id = str(
                        child.get("@id") or child.get("id") or child.get("name") or ""
                    ).strip()
                    if target_id:
                        payload.relations.append(
                            RelationRecord(source=source, relation=str(key), target=target_id)
                        )
                    self._json_relations(child, payload=payload, parent=f"{source}.{key}")
                else:
                    self._json_relations(child, payload=payload, parent=f"{source}.{key}")
        elif isinstance(value, list):
            for index, child in enumerate(value):
                self._json_relations(child, payload=payload, parent=f"{parent}[{index}]")

    def _read_pdb(self, file_path: Path, payload: ParsedAttachmentPayload) -> None:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        payload.text_blocks.append(TextBlock(text=text))
        for line in text.splitlines():
            if not line.startswith(("ATOM  ", "HETATM")):
                continue
            try:
                payload.coordinates.append(
                    CoordinateRecord(
                        identifier=line[6:11].strip() or str(len(payload.coordinates) + 1),
                        x=float(line[30:38]),
                        y=float(line[38:46]),
                        z=float(line[46:54]),
                        attributes={
                            "record": line[:6].strip(),
                            "atom": line[12:16].strip(),
                            "residue": line[17:20].strip(),
                            "chain": line[21:22].strip(),
                            "residue_index": line[22:26].strip(),
                        },
                    )
                )
            except (TypeError, ValueError):
                continue

    def _read_page_blocks(
        self,
        content: str,
        payload: ParsedAttachmentPayload,
        *,
        marker: str,
    ) -> None:
        pattern = re.compile(rf"\[{marker}\s+(\d+)\]\s*", flags=re.IGNORECASE)
        matches = list(pattern.finditer(content or ""))
        if not matches:
            if content.strip():
                payload.text_blocks.append(TextBlock(text=content.strip()))
            return
        for index, match in enumerate(matches):
            end = matches[index + 1].start() if index + 1 < len(matches) else len(content)
            text = content[match.end():end].strip()
            if text:
                payload.text_blocks.append(
                    TextBlock(
                        text=text,
                        page=int(match.group(1)),
                        block_type=marker,
                    )
                )

    def _read_visual_content(
        self,
        content: str,
        payload: ParsedAttachmentPayload,
    ) -> None:
        text = str(content or "").strip()
        if text.startswith("Ollama vision model:"):
            text = text.split("\n", 1)[1] if "\n" in text else ""
        data = self._json_object(text)
        if not isinstance(data, dict):
            if text:
                payload.visual_blocks.append(VisualBlock(text=text, object_type="vision_text"))
            return
        ocr_blocks = data.get("ocr_blocks") or []
        for index, block in enumerate(ocr_blocks):
            if isinstance(block, dict):
                block_text = str(block.get("text") or "").strip()
                region = str(block.get("region") or block.get("position") or "")
                attributes = {
                    key: value for key, value in block.items() if key not in {"text", "region", "position"}
                }
            else:
                block_text = str(block).strip()
                region = ""
                attributes = {}
            if block_text:
                payload.visual_blocks.append(
                    VisualBlock(
                        text=block_text,
                        region=region,
                        object_type="ocr",
                        attributes={"index": index, **attributes},
                    )
                )
        summary_attributes = {
            key: data.get(key)
            for key in (
                "objects",
                "numbers",
                "colors",
                "spatial_relations",
                "uncertainties",
                "grid",
                "candidate_words",
            )
            if data.get(key) not in (None, "", [], {})
        }
        payload.visual_blocks.append(
            VisualBlock(
                text=str(data.get("summary") or "").strip(),
                object_type="vision_summary",
                attributes=summary_attributes,
            )
        )

    @staticmethod
    def _json_object(text: str) -> dict[str, Any] | None:
        try:
            value = json.loads(text)
            return value if isinstance(value, dict) else None
        except Exception:
            start = text.find("{")
            end = text.rfind("}")
            if start < 0 or end <= start:
                return None
            try:
                value = json.loads(text[start:end + 1])
                return value if isinstance(value, dict) else None
            except Exception:
                return None

    @staticmethod
    def _cell_metadata(cell: Any) -> dict[str, Any]:
        metadata: dict[str, Any] = {}
        fill = getattr(cell, "fill", None)
        color = getattr(fill, "fgColor", None)
        color_value = str(
            getattr(color, "rgb", "")
            or getattr(color, "indexed", "")
            or getattr(color, "theme", "")
            or ""
        ).strip()
        if color_value and color_value not in {"00000000", "0"}:
            metadata["fill_color"] = color_value
        number_format = str(getattr(cell, "number_format", "") or "").strip()
        if number_format and number_format != "General":
            metadata["number_format"] = number_format
        return metadata


__all__ = ["AttachmentPayloadBuilder"]
